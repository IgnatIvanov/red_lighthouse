import os
import json
import string
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import Font
from openpyxl import load_workbook
from openpyxl.styles import Alignment

from docx import Document
from docx.shared import Mm  # Для установки значений интервалов в миллиметрах
from docxtpl import DocxTemplate
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_LINE_SPACING  # Для установки междустрочных интервалов
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from .models import Dog
from .models import Breed
from .models import DogClass
from .models import Participant



BASE_DIR = Path(__file__).resolve().parent.parent
templates_path = BASE_DIR / 'out_doc/templates/out_doc/'



# Преобразование даты к строке с сохранение лидирующих нолей
def toDateStr(number):


    str_number = str(number)
    if len(str_number) > 3:
        # Если передан год для преобразования
        return str_number
    if len(str_number) < 2:
        # Передано однозначное число дня или месяца
        str_number = '0' + str_number

    return str_number



# Функция печати временных сертификатов
def print_temp_sertif(events, temp_path, project_name):

    # Загрузка документа
    template_name = 'временный сертификат.docx'
    template_file = templates_path / template_name

    for event in events:
        
        save_path = temp_path / project_name / ('Тестирование ' + event['rank'] + ' ' + event['comment'])
    
        if not os.path.exists(BASE_DIR / save_path):
            os.makedirs(BASE_DIR / save_path)  # Создание пути сохранения файла

        dogs_list = event['participants_data']
        for i in range(len(dogs_list)):

            doc = DocxTemplate(template_file)

            dogie = dogs_list[i]

            sex = '{} \ {}'.format(dogie['sex_ru'], dogie['sex_en'])

            # Подстановка данных
            context = {
                'name': dogie['name'],
                'sex': sex,
                'tattoo': dogie['tattoo'],
                'rkf': dogie['rkf'],
                'birth': dogie['birth_date'],
                'owner': dogie['owner'],
                'breed': dogie['breeder']
            }
            doc.render(context)

            # Заменяем в клейме все знаки препинания на точки
            # save_file_name = dogs_list[i]['tattoo'] + '.docx'
            punct = string.punctuation
            punct = punct.replace('.', '')
            save_file_name = dogie['tattoo'].translate(str.maketrans('', '', punct)) + '.docx'

            # Сохранение документа
            save_file = save_path / save_file_name
            doc.save(save_file)

    return



# Создание каталогов для каждого события отдельно
def create_events_catalogs(events, temp_path, project_name):


    # Загрузка документа
    template_name = 'оглавление.docx'
    template_file = templates_path / template_name


    for event in events:

        # Создание объекта документа
        document = Document(template_file)

        # Применение форматирования ко всему документу
        style = document.styles['Normal']
        style.font.name = 'Times New Roman'


        dogs_list = event['participants_data']

        # ===========================================================================
        # Работа с оглавлением

        table = document.tables[0]

        header_lines = []
        start_npp = {}
        end_npp = {}

        # Сбор данных для наполнения каталога
        for i in range(len(dogs_list)):

            dogie = dogs_list[i]
            npp = dogie['npp']
            dog_id = dogs_list[i]['dog_id']
            dog_obj = Dog.objects.filter(id=dog_id).first()
            breed_obj = Breed.objects.filter(id=dog_obj.breed_id).first()

            breed_line = '{} ({}) \ {} ({})'.format(
                breed_obj.name_ru,
                breed_obj.country_ru,
                breed_obj.name_en,
                breed_obj.country_en
            )

            line = {
                'fci': dogie['fci'],
                'breed_id': breed_obj.bid,
                'breed': breed_line,
                'count': 1,
                'start_npp': npp,
                'end_npp': npp,
                'npp': npp,
            }
            header_lines.append(line)


        # Уплотнение
        for i in range(len(header_lines)-1, 0, -1):
            current_line = header_lines[i]
            prev_line = header_lines[i - 1]
            if current_line['breed_id'] == prev_line['breed_id']:
                header_lines[i - 1]['count'] += 1
                header_lines[i - 1]['end_npp'] = header_lines[i]['end_npp']
                header_lines.pop(i)


        # Добавление тела таблицы
        current_fci = -1
        for i in range(len(header_lines)):

            line = header_lines[i]
            fci = line['fci']
            breed = line['breed']
            breed_id = line['breed_id']
            count = line['count']
            start_npp = line['start_npp']
            end_npp = line['end_npp']
            
            
            # Добавление заголовка группы FCI
            if fci != current_fci:
                row = table.add_row()
                current_fci = fci
                value = str(current_fci) + ' ГРУППА F.C.I.'
                p = row.cells[0].paragraphs[0]
                run = p.add_run(value)
                run.font.bold = True
                run.font.size = Pt(14)
                p_fmt = p.paragraph_format
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[0].merge(row.cells[3])

            # Добавление записи
            # Добавление новой строки
            row = table.add_row()

            # Добавление кода породы
            p = row.cells[0].paragraphs[0]
            value = str(breed_id)
            run = p.add_run(value)
            run.font.size = Pt(12)
            p_fmt = p.paragraph_format        
            p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Добавление названия породы
            p = row.cells[1].paragraphs[0]
            run = p.add_run(breed)
            run.font.size = Pt(12)

            # Добавление количества собак
            p = row.cells[2].paragraphs[0]
            value = str(count)
            run = p.add_run(value)
            run.font.size = Pt(12)
            p_fmt = p.paragraph_format        
            p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Добавления номеров начала и конца
            if start_npp == end_npp:
                p = row.cells[3].paragraphs[0]
                value = str(start_npp)
                run = p.add_run(value)
                run.font.size = Pt(12)
                p_fmt = p.paragraph_format        
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            else:
                p = row.cells[3].paragraphs[0]
                value = '{} - {}'.format(start_npp, end_npp)
                run = p.add_run(value)
                run.font.size = Pt(12)
                p_fmt = p.paragraph_format        
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Конец работы с оглавлением
        # ===========================================================================


        current_fci = ''
        current_breed = ''
        current_sex = ''
        current_class = ''
        current_line = 1

        for i in range(len(dogs_list)):
            fci = dogs_list[i]['fci']
            dog_id = dogs_list[i]['dog_id']
            dog_obj = Dog.objects.filter(id=dog_id).first()
            participant_id = dogs_list[i]['participant_id']
            breed_obj = Breed.objects.filter(id=dog_obj.breed_id).first()
            parts_object = Participant.objects.filter(id=participant_id).first()
            class_obj = DogClass.objects.filter(id=parts_object.class_id).first()
            
            # Добавление заголовка группы FCI
            if fci != current_fci:

                # Вставка группы FCI в документ
                current_fci = fci
                p = document.add_paragraph()
                value = str(current_fci) + ' ГРУППА F.C.I.'
                run = p.add_run(value)

                # Форматирование группы FCI
                run.font.bold = True
                run.font.size = Pt(14)
                p_fmt = p.paragraph_format
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_fmt.space_before = Pt(6)
                p_fmt.space_after = Pt(12)

            name_ru = breed_obj.name_ru
            country_ru = breed_obj.country_ru
            name_en = breed_obj.name_en
            country_en = breed_obj.country_en
            breed_str = '{} ({}) \ {} ({})'.format(name_ru, country_ru, name_en, country_en)

            # Добавление раздела породы 
            if breed_str != current_breed:
                current_breed = breed_str
                p = document.add_paragraph()
                run = p.add_run(current_breed)

                # форматирование раздела породы
                run.font.bold = True
                run.font.size = Pt(12)
                p_fmt = p.paragraph_format
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_fmt.space_before = Pt(0)
                p_fmt.space_after = Pt(0)

                # Добавление раздела судьи и ринга

                # Получение данных 
                # Загрузка записей назначенных судей и рингов
                judges_json_path = BASE_DIR / 'judges.json'
                with open(judges_json_path, 'r', encoding='utf8') as judges_file:
                    judges_json = json.load(judges_file)

                # Чтение назначенных судьи и ринга
                current_ring = 'Не назначен'
                current_judge = 'Не назначен'
                judges_id = str(event['id']) + '-' + str(dog_obj.breed_id)
                for el in judges_json:
                    if el['judges_id'] == judges_id:
                        current_ring = el['ring']
                        current_judge = el['judge']

                judge_ring_str = 'Судья - {}, Ринг {}'.format(
                    current_judge,
                    current_ring
                )
                p = document.add_paragraph()
                run = p.add_run(judge_ring_str)
                run.font.size = Pt(12)
                p_fmt = p.paragraph_format
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_fmt.space_before = Pt(0)
                p_fmt.space_after = Pt(0)


            sex_str = 'СУКИ \ FEMALES'
            if dog_obj.is_male == True:
                sex_str = 'КОБЕЛИ \ MALES'

            if sex_str != current_sex:
                current_sex = sex_str
                p = document.add_paragraph()
                run = p.add_run(current_sex)

                # форматирование раздела пола
                run.font.bold = True
                run.font.italic = True
                run.font.size = Pt(12)
                p_fmt = p.paragraph_format
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_fmt.space_before = Pt(0)
                p_fmt.space_after = Pt(0)

            # Добавление раздела класса
            class_str = 'Класс: {} \ {} class'.format(class_obj.name_ru, class_obj.name_en.capitalize())

            if class_str != current_class:
                current_class = class_str
                p = document.add_paragraph()
                run = p.add_run(current_class)

                # форматирование раздела класса
                run.font.bold = True
                run.font.underline = True
                run.font.size = Pt(12)
                p_fmt = p.paragraph_format
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_fmt.space_before = Pt(0)
                p_fmt.space_after = Pt(0)

            # Создаём объект абзаца собаки
            dogie_paragragh = document.add_paragraph()

            # Добавление номера по порядку в каталоге
            run = dogie_paragragh.add_run(str(dogs_list[i]['npp']) + '. ')
            run.font.bold = True
            run.font.size = Pt(12)

            # Добавление клички
            run = dogie_paragragh.add_run(dogs_list[i]['name'] + ', ')
            run.font.bold = True
            run.font.size = Pt(12)

            # Добавление РКФ
            run = dogie_paragragh.add_run('РКФ ' + dogs_list[i]['rkf'] + ', ')
            run.font.size = Pt(12)
            
            # Добавление даты рождения
            run = dogie_paragragh.add_run('д.р. ' + dogs_list[i]['birth_date'] + ' г., ')
            run.font.size = Pt(12)

            # Добавление клейма
            run = dogie_paragragh.add_run(dogs_list[i]['tattoo'] + ', ')
            run.font.size = Pt(12)

            # Добавление окраски собаки
            colour_str = dog_obj.colour_ru
            if colour_str == '-':
                colour_str = dog_obj.colour_en
            run = dogie_paragragh.add_run(colour_str + ', ')
            run.font.size = Pt(12)

            # Добавление родителей собаки
            family_str = '({} X {}), '.format(
                dogs_list[i]['father_name'],
                dogs_list[i]['mother_name']
            )
            run = dogie_paragragh.add_run(family_str)

            # Разобраться с вводом родителей собаки

            # father_obj = Dog.objects.filter(tattoo=dog_obj.father_tattoo).first()
            # mother_obj = Dog.objects.filter(tattoo=dog_obj.mother_tattoo).first()

            # father_name = father_obj.name_ru
            # if father_name == '-':
            #     father_name = father_obj.name_en

            # mother_name = mother_obj.name_ru
            # if mother_name == '-':
            #     mother_name = mother_obj.name_en

            # dogie_str += father_name + 'X' + mother_name + ', '

            # Добавление заводчика собаки
            run = dogie_paragragh.add_run('зав. ' + dog_obj.breeder + ', ')
            run.font.size = Pt(12)

            # Добавление владельца собаки
            run = dogie_paragragh.add_run('вл. ' + dog_obj.owner + ', ')
            run.font.size = Pt(12)

            # Добавление кортокого адреса            
            run = dogie_paragragh.add_run(dogs_list[i]['short_address'] + '.')
            run.font.size = Pt(12)


        # Путь сохранения нового каталога
        save_path = temp_path / project_name
        
        if not os.path.exists(save_path):
            os.makedirs(save_path)  # Создание пути сохранения файла

        save_path = save_path / ('Каталог ' + event['rank'] + ' ' + event['comment'] + '.docx')

        document.save(save_path)
        del document


    return



# Создание отчётов на каждое событие
def create_events_reports(events, temp_path, project_name):
    
    # print('Создание отчётов')b

    for event in events:

        # Шаблон для САС ЧРКФ, САС ЧФ РФЛС
        # Будет также установлен для КЧК и КЧП
        template_name = 'итоговый отчёт САС ЧФ РФЛС, САС ЧРКФ.xlsx'
        date_table_address = 'V11'

        # Изменение шаблона для монопородных выставок
        if event['type'] == 'Монопородные':
            template_name = 'итоговый отчёт монопородный.xlsx'
            date_table_address = 'O11'
            # template_file = templates_path / template_name
            # print('template_file', template_file)
            # wb2 = load_workbook(template_file)
            # print(wb2.sheetnames)
            # print(event['type'], event['rank'])

        # Изменение шаблона для выставок типа САС группы
        if event['rank'].endswith('гр.'):
            template_name = 'итоговый отчёт САС группы.xlsx'
            date_table_address = 'AB11'
            # template_file = templates_path / template_name
            # print(event['type'], event['rank'])


        # print(event['type'], event['rank'])
        # Загружаем файл шаблона
        template_file = templates_path / template_name
        wb = load_workbook(template_file)
        # Делаем единственный лист активным
        ws = wb.active

        # Изменение заголовка отчёта
        # Перезапись даты события в заголовке
        event_date = toDateStr(event['date'].day) + '.' + toDateStr(event['date'].month) + '.' + toDateStr(event['date'].year)
        ws['E6'] = event_date
        # Перезапись даты в первой строке таблицы
        ws[date_table_address] = event_date
        # Перезапись ранга события в заголовке
        rank = event['rank'].replace(' гр.', ' группы')
        ws['E8'] = rank

        # Запись событий
        # Если на событие никто не записался
        current_str = 11
        if not len(event['participants_data']):
            # Удалим строку с шаблоном форматирования 
            ws.delete_cols(11, 1)
        else:
            # Если запись на событие есть
            # То записываем первую строку в таблицу
            # Остальные строки будем записывать в цикле
            participants_data = event['participants_data']
            first_dogie = participants_data[0]
            ws['A' + str(current_str)] = first_dogie['breed_ru']
            ws['B' + str(current_str)] = 'Судья не прочитан'
            ws['C' + str(current_str)] = first_dogie['npp']
            ws['D' + str(current_str)] = first_dogie['name']
            ws['E' + str(current_str)] = first_dogie['birth_date']
            ws['F' + str(current_str)] = first_dogie['rkf']
            ws['G' + str(current_str)] = first_dogie['class_ru']
            del participants_data[0]
            print(ws.row_dimensions[11])
            current_str += 1


        # Записываем остальных участников в таблицу
        for dogie in participants_data:
            # pass
            # Копирование предыдущей строки
            # for cur_row in range(50):
            #     try:
            #         # Пробуем скопировать ячейку предыдущей строки
            #         # ws[current_str][row].value = ws[current_str - 1][row].value
            #         # ws[current_str][row].fill = ws[current_str - 1][row].fill
            #         ws.cell(row=cur_row, column=current_str).fill = ws.cell(row=cur_row, column=current_str - 1).fill
                    
            #     except Exception as err:
            #         # Пропускаем, если ячейка не заполнена
            #         # print("An exception occurred")
            #         print('Can`t copy format from previous string')
            #         print(err)
            #         continue
                # ws[current_str][row].value = ws[current_str - 1][row].value
            ws['A' + str(current_str)] = dogie['breed_ru']
            ws['B' + str(current_str)] = 'Судья не прочитан'
            ws['C' + str(current_str)] = dogie['npp']
            ws['D' + str(current_str)] = dogie['name']
            ws['E' + str(current_str)] = dogie['birth_date']
            ws['F' + str(current_str)] = dogie['rkf']
            ws['G' + str(current_str)] = dogie['class_ru']
            # ws.rows[current_str].value = ws.rows[current_str - 1].value

            # for row in ws.rows:
            #     print(row[0].value)
            current_str += 1












        # Директория сохранения нового отчёта
        save_path = temp_path / project_name
        
        if not os.path.exists(save_path):
            os.makedirs(save_path)  # Создание пути сохранения файла

        # Полный путь сохранения файла отчёта
        save_path = save_path / ('Отчёт ' + event['rank'] + ' ' + event['comment'] + '.xlsx')
        
        # Сохранение текущего отчёта
        wb.save(save_path)
        del wb

        continue



        

        template_name = 'временный сертификат.docx'
        template_file = templates_path / template_name

        # Путь сохранения нового отчёта
        save_path = temp_path / project_name
        
        if not os.path.exists(save_path):
            os.makedirs(save_path)  # Создание пути сохранения файла

        save_path = save_path / ('Отчёт ' + event['rank'] + ' ' + event['comment'] + '.xlsx')

        # Создание чистого Excel документа
        wb = Workbook()

        # Делаем единственный лист активным
        ws = wb.active

        # Вставка заголовка отчёта
        ws['B2'] = 'ИТОГОВЫЙ ОТЧЕТ'

        ws['B4'] = 'Название кинологической организации'
        ws['E4'] = 'МЕЖРЕГИОНАЛЬНАЯ ОБЩЕСТВЕННАЯ ОРГАНИЗАЦИЯ КЛУБ ПЛЕМЕННОГО СОБАКОВОДСТВА "КРАСНЫЙ МАЯК"'

        ws['B5'] = 'Название выставки'
        ws['E5'] = 'СЕРТИФИКАТНАЯ ВЫСТАВКА "КРАСНЫЙ МАЯК"'

        ws['B6'] = 'Дата проведения'
        event_date = event['date']
        ws['E6'] = str(event_date)

        ws['B7'] = 'Город'
        ws['E7'] = 'МОСКВА'

        ws['B8'] = 'Ранг выставки'
        ws['E8'] = event['rank']

        # Вставка шапки таблицы



        # Сохранение текущего отчёта
        wb.save(save_path)
        del wb



# Создание закрытого списка участников для сверки
def create_private_list(events, temp_path, project_name):

    rows = []
    
    # Сбор строк
    for event in events:        

        # Сбор данных о событии
        event_date = event['date']
        event_field = str()
        event_field += event['org'] + ', '
        event_field += event['type'] + ', '
        event_field += event['rank'] + ', '
        event_field += toDateStr(event_date.day) + '.' + toDateStr(event_date.month) + '.' + toDateStr(event_date.year) + ', ' 
        event_field += event['comment']

        # Добавление новой строки в список
        for dogie in event['participants_data']:
            new_row = {}
            new_row['event'] = event_field
            new_row['breed'] = dogie['breed_ru']
            new_row['class'] = dogie['class_ru']
            new_row['sex'] = dogie['sex_ru']
            new_row['tattoo'] = dogie['tattoo']
            new_row['owner'] = dogie['owner']
            rows.append(new_row)

    
    # Сортировка собак по клейму 
    rows = sorted(rows, key=lambda x: x['tattoo'])
        
    # Путь сохранения нового отчёта
    save_path = temp_path / project_name
    
    if not os.path.exists(save_path):
        os.makedirs(save_path)  # Создание пути сохранения файла

    save_path = save_path / ('Список закрытый для сверки.xlsx')

    # Создание чистого Excel документа
    wb = Workbook()

    # Делаем единственный лист активным
    ws = wb.active

    # Вставка заголовка списка
    ws['A1'] = 'Событие'
    ws['B1'] = 'Порода'
    ws['C1'] = 'Класс'
    ws['D1'] = 'Пол'
    ws['E1'] = 'Клеймо'
    ws['F1'] = 'Владелец'
    current_line = 2

    # Для настройки ширины столбцов
    a_width = 0
    b_width = 0
    c_width = 0
    d_width = 0
    e_width = 0
    f_width = 0

    # Вставка данных
    for row in rows:
        # print(row)
        ws['A' + str(current_line)] = row['event']
        ws['B' + str(current_line)] = row['breed']
        ws['C' + str(current_line)] = row['class']
        ws['D' + str(current_line)] = row['sex']
        ws['E' + str(current_line)] = row['tattoo']
        ws['F' + str(current_line)] = row['owner']
        # ws.row_dimensions[current_line].width = True
        current_line += 1
        
        # Сохранение максимальной ширины столбцов 
        a_width = max(a_width, len(row['event']))
        b_width = max(b_width, len(row['breed']))
        c_width = max(c_width, len(row['class']))
        d_width = max(d_width, len(row['sex']))
        e_width = max(e_width, len(row['tattoo']))
        f_width = max(f_width, len(row['owner']))


    # Настройка ширины столбцов
    ws.column_dimensions['A'].width = a_width
    ws.column_dimensions['B'].width = b_width
    ws.column_dimensions['C'].width = c_width
    ws.column_dimensions['D'].width = d_width
    ws.column_dimensions['E'].width = e_width
    ws.column_dimensions['F'].width = f_width

    # Сохранение текущего отчёта
    wb.save(save_path)
    # print('save_path', save_path)
    del wb



# Создание открытого списка участников для сверки
def create_open_list(events, temp_path, project_name):


    document = Document()
    
    # Сбор строк
    for event in events:              

        rows = []

        # Сбор данных о событии
        event_date = event['date']
        event_field = str()
        event_field += event['org'] + ', '
        event_field += event['type'] + ', '
        event_field += event['rank'] + ', '
        event_field += toDateStr(event_date.day) + '.' + toDateStr(event_date.month) + '.' + toDateStr(event_date.year) + ', ' 
        event_field += event['comment']

        # Запись заголовка с названием события
        p = document.add_paragraph('')
        p.add_run(event_field.upper() + '\n').bold = True  # Добавление жирного выделения
        fmt = p.paragraph_format  # Начинаем форматирование заголовка (название события)
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE  # Устанавливаем одинарный междустрочный интервал
        fmt.space_before = Mm(0)  # Установка отступа перед абзацем
        fmt.space_after = Mm(0)  # Установка отступа после абзаца

        # Добавление новой строки в список
        for dogie in event['participants_data']:
            # new_row = {}
            # new_row['event'] = event_field
            # Добавление данных о породе в абзац
            breed_data = dogie['breed_ru']
            p = document.add_paragraph(breed_data + '; ')
            # new_row['breed'] = '{} \ {}'.format(dogie['breed_ru'], dogie['breed_en'])

            # Добавление данных о классе в абзац
            class_data = dogie['class_ru']
            p.add_run(class_data + '; ')
            # new_row['class'] = '{} \ {}'.format(dogie['class_ru'], dogie['class_en'])

            # Добавление данных о поле в абзац
            sex_data = dogie['sex_ru']
            p.add_run(sex_data + '; ')
            # new_row['sex'] = '{} \ {}'.format(dogie['sex_ru'], dogie['sex_en'])

            # Добавление данных о клейме в абзац
            p.add_run(dogie['tattoo'])
            # new_row['tattoo'] = dogie['tattoo']
            # new_row['owner'] = dogie['owner']
            # rows.append(new_row)
            # print(new_row, '\n')

            # Форматирование строки данных о собаке
            fmt = p.paragraph_format  # Начинаем форматирование
            fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE  # Устанавливаем одинарный междустрочный интервал
            fmt.space_before = Mm(0)
            fmt.space_after = Mm(0)

        
        # Добавление пустой строки после перечисления всех собак
        p = document.add_paragraph('')
        fmt = p.paragraph_format  # Начинаем форматирование
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE  # Устанавливаем одинарный междустрочный интервал
        fmt.space_before = Mm(0)
        fmt.space_after = Mm(0)

    
        # Сортировка собак по клейму 
        # rows = sorted(rows, key=lambda x: x['tattoo'])
        
    # Путь сохранения нового отчёта
    save_path = temp_path / project_name
    
    if not os.path.exists(save_path):
        os.makedirs(save_path)  # Создание пути сохранения файла

    save_path = save_path / ('Список открытый для сверки.docx')

    document.save(save_path)
    del document



# Создание дипломов
def create_diploms(events, temp_path, project_name):
    

    # Загрузка документа
    template_name = 'диплом.docx'
    template_file = templates_path / template_name

    for event in events:

        save_path = temp_path / project_name / ('Дипломы ' + event['rank'] + ' ' + event['comment'])
    
        if not os.path.exists(BASE_DIR / save_path):
            os.makedirs(BASE_DIR / save_path)  # Создание пути сохранения файла


        event_type = event['type']
        event_rank = event['rank']
        event_comment = event['comment']
        event_date = '{}.{}.{}'.format(
            toDateStr(event['date'].day), 
            toDateStr(event['date'].month), 
            toDateStr(event['date'].year),
        )

        event_name = str()
        if event_type == 'Монопородная':
            event_name = '{} {}'.format(
                event_type,
                event_comment
            ).upper()
        else:
            event_name = '{} выставка {}'.format(
                event_type,
                event_rank
            ).upper()

        dogs_list = event['participants_data']
        for i in range(len(dogs_list)):

            dogie = dogs_list[i]
            dog_id = dogs_list[i]['dog_id']
            dog_obj = Dog.objects.filter(id=dog_id).first()
            breed_obj = Breed.objects.filter(id=dog_obj.breed_id).first()

            # Подстановка данных
            context = {
                'event_name': event_name,
                'event_date': event_date,
                'breed_name': dogie['breed_ru'],
                'breed_country': breed_obj.country_ru,
                'dog_name': dogie['name'],
                'npp': dogie['npp'],
                'sex': dogie['sex_ru'].lower(),
                'class': dogie['class_ru'],
                'owner': dogie['owner'],
                'colour': dog_obj.colour_ru,
            }
            doc = DocxTemplate(template_file)
            doc.render(context)

            # Заменяем в клейме все знаки препинания на точки
            punct = string.punctuation
            punct = punct.replace('.', '')
            save_file_name = dogie['tattoo'].translate(str.maketrans('', '', punct)) + '.docx'

            # Сохранение документа
            save_file = save_path / save_file_name
            doc.save(save_file)


    return